#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取Word文档中的剧本简介和角色简介内容
"""

import os
import re
from docx import Document
import json

def extract_docx_content(docx_path):
    """提取Word文档内容"""
    print(f"\n=== 提取文件: {os.path.basename(docx_path)} ===")
    
    try:
        doc = Document(docx_path)
        content = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # 尝试获取表格内容
        tables_content = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    tables_content.append(" | ".join(row_text))
        
        # 判断文档类型（剧本简介或角色简介）
        is_script_intro = False
        is_role_intro = False
        
        for text in content[:10]:  # 检查前10段
            text_lower = text.lower()
            if "剧本" in text_lower or "简介" in text_lower or "介绍" in text_lower:
                is_script_intro = True
            if "角色" in text_lower or "人物" in text_lower or "玩家" in text_lower or "人设" in text_lower:
                is_role_intro = True
        
        print(f"内容分析:")
        print(f"- 段落数量: {len(content)}")
        print(f"- 表格行数: {len(tables_content)}")
        print(f"- 疑似剧本简介: {is_script_intro}")
        print(f"- 疑似角色简介: {is_role_intro}")
        
        # 显示部分内容
        print("\n前20个段落:")
        for i, text in enumerate(content[:20]):
            print(f"{i+1:2}. {text[:80]}{'...' if len(text) > 80 else ''}")
        
        if tables_content:
            print(f"\n表格内容（前10行）:")
            for i, row in enumerate(tables_content[:10]):
                print(f"{i+1:2}. {row[:100]}{'...' if len(row) > 100 else ''}")
        
        return {
            "content": content,
            "tables": tables_content,
            "is_script_intro": is_script_intro,
            "is_role_intro": is_role_intro,
            "total_paragraphs": len(content),
            "total_table_rows": len(tables_content)
        }
        
    except Exception as e:
        print(f"读取失败: {e}")
        return None

def analyze_for_update(script_name, content_data):
    """分析内容如何更新网站"""
    print(f"\n=== 分析 {script_name} 更新需求 ===")
    
    content = content_data.get("content", [])
    tables = content_data.get("tables", [])
    
    # 查找可能的剧本简介
    script_intro = ""
    for i, para in enumerate(content):
        if len(para) > 30 and ("简介" in para or "剧本" in para):
            # 获取接下来的几个段落作为剧本简介
            script_intro = para
            if i+1 < len(content) and len(content[i+1]) > 10:
                script_intro += " " + content[i+1]
            if i+2 < len(content) and len(content[i+2]) > 10:
                script_intro += " " + content[i+2]
            print(f"找到剧本简介（长度: {len(script_intro)}）")
            break
    
    # 查找角色数据
    roles = []
    role_section_found = False
    
    for para in content:
        para_lower = para.lower()
        if "角色" in para_lower or "人物" in para_lower or "玩家" in para_lower:
            role_section_found = True
            print(f"找到角色章节: {para[:50]}...")
            continue
        
        # 尝试匹配角色格式：名字（性别）描述
        # 例如：李靖（男）大唐名将...
        role_patterns = [
            r'^([^（]+)（([男女MF])）(.+)$',
            r'^([^:：]+)[:：]\s*(.+)$',
            r'^\d+\.\s*([^（]+)（([男女MF])）(.+)$',
        ]
        
        for pattern in role_patterns:
            match = re.match(pattern, para)
            if match:
                if len(match.groups()) == 3:
                    name, gender, desc = match.groups()
                    gender_map = {"男": "male", "女": "female", "M": "male", "F": "female"}
                    roles.append({
                        "name": name.strip(),
                        "gender": gender_map.get(gender, "male"),
                        "desc": desc.strip()
                    })
                elif len(match.groups()) == 2:
                    name, desc = match.groups()
                    # 默认男性，因为没有指定性别
                    roles.append({
                        "name": name.strip(),
                        "gender": "male",
                        "desc": desc.strip()
                    })
    
    # 从表格中提取角色信息
    if tables and not roles:
        for row in tables:
            # 表格格式可能是：角色名 | 性别 | 简介
            parts = row.split("|")
            if len(parts) >= 3:
                name = parts[0].strip()
                gender = parts[1].strip()
                desc = parts[2].strip()
                gender_map = {"男": "male", "女": "female", "男性": "male", "女性": "female"}
                roles.append({
                    "name": name,
                    "gender": gender_map.get(gender, "male"),
                    "desc": desc
                })
            elif len(parts) >= 2:
                # 可能没有性别列
                name = parts[0].strip()
                desc = parts[1].strip()
                roles.append({
                    "name": name,
                    "gender": "male",  # 默认男性
                    "desc": desc
                })
    
    print(f"发现角色数量: {len(roles)}")
    for i, role in enumerate(roles[:5]):
        print(f"  {i+1}. {role['name']} ({role['gender']}): {role['desc'][:50]}{'...' if len(role['desc']) > 50 else ''}")
    
    return {
        "script_name": script_name,
        "script_intro": script_intro if script_intro else ("这是" + script_name + "的剧本简介"),
        "roles": roles,
        "has_roles": len(roles) > 0,
        "has_intro": bool(script_intro)
    }

def main():
    """主函数"""
    # 定义剧本文件映射
    script_files = {
        "请君入瓮": "C:/Users/admin/Desktop/请君入瓮.docx",
        "万仞青山听水寒": "C:/Users/admin/Desktop/万仞青山听水寒.docx"
    }
    
    all_data = {}
    
    for script_name, file_path in script_files.items():
        if os.path.exists(file_path):
            print(f"\n{'='*60}")
            print(f"处理剧本: {script_name}")
            print(f"文件路径: {file_path}")
            
            # 提取内容
            content_data = extract_docx_content(file_path)
            if content_data:
                # 分析如何更新网站
                update_info = analyze_for_update(script_name, content_data)
                all_data[script_name] = update_info
        else:
            print(f"\n文件不存在: {file_path}")
    
    # 输出摘要
    print(f"\n{'='*60}")
    print("提取摘要:")
    print("="*60)
    
    for script_name, data in all_data.items():
        print(f"\n{script_name}:")
        print(f"  - 有剧本简介: {data['has_intro']}")
        print(f"  - 角色数量: {len(data['roles'])}")
        print(f"  - 是否需要更新: {'是' if data['has_intro'] or data['has_roles'] else '否'}")
        
        if data['has_intro']:
            print(f"  - 简介预览: {data['script_intro'][:80]}...")
        
        if data['roles']:
            print(f"  - 角色示例: {data['roles'][0]['name']} ({data['roles'][0]['gender']})")
    
    # 保存提取结果
    output_file = "extracted_script_data.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n提取结果已保存到: {output_file}")
    
    return all_data

if __name__ == "__main__":
    main()