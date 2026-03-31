#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析所有剧本的简介和角色信息，生成更新数据
"""

import os
import re
import json
from docx import Document

def extract_script_descriptions(file_path):
    """从请君入瓮.docx提取所有剧本简介"""
    print(f"\n=== 从 {os.path.basename(file_path)} 提取剧本简介 ===")
    
    try:
        doc = Document(file_path)
        content = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        scripts = {}  # 剧本名 -> 简介
        current_script = None
        
        for para in content:
            # 检查是否是剧本标题行（包含冒号或剧本名特征）
            if "：" in para or ":" in para or len(para) < 10:
                script_name = para.replace("：", ":").split(":")[0].strip()
                if script_name and len(script_name) > 1 and script_name not in ["请君入瓮", "流芳", "如故", "阙落", "万仞青山听水寒", "梦灵祈念时", "窃云台", "壁上观", "同归", "剧本"]:
                    # 这可能是真正的剧本名（但不是我们已知的完整名单）
                    pass
                else:
                    # 尝试匹配已知剧本
                    for known_script in ["请君入瓮", "流芳", "如故", "阙落", "万仞青山听水寒", "梦灵祈念时", "窃云台", "壁上观", "同归"]:
                        if known_script in para:
                            script_name = known_script
                            # 提取简介部分
                            desc_part = para.split(script_name + "：")[-1].split(script_name + ":")[-1]
                            if desc_part.strip():
                                scripts[script_name] = {"description": desc_part.strip()}
                            current_script = script_name
                            break
        
        # 查看前20行内容进行分析
        print("\n文件前20行内容:")
        for i, para in enumerate(content[:20]):
            print(f"{i+1:2}. {para}")
        
        # 基于文件内容手动提取（因为格式特殊）
        if "请君入瓮：" in content[0]:
            # 格式是：剧本名：简介
            for i in range(0, len(content), 2):
                if i+1 < len(content):
                    line = content[i]
                    for script_name in ["请君入瓮", "流芳", "如故", "阙落", "万仞青山听水寒", "梦灵祈念时", "窃云台", "壁上观", "同归"]:
                        if line.startswith(script_name + "："):
                            desc_line = content[i+1] if i+1 < len(content) else ""
                            scripts[script_name] = {"description": desc_line.strip()}
                            print(f"找到 {script_name}: {desc_line[:60]}...")
        
        print(f"\n提取结果: {len(scripts)} 个剧本")
        for name, data in scripts.items():
            print(f"  - {name}: {data['description'][:80]}...")
        
        return scripts
        
    except Exception as e:
        print(f"提取失败: {e}")
        return {}

def extract_wanren_roles(file_path):
    """从万仞青山听水寒.docx提取角色信息"""
    print(f"\n=== 从 {os.path.basename(file_path)} 提取角色信息 ===")
    
    try:
        doc = Document(file_path)
        content = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        roles = []
        current_role = None
        
        for para in content:
            # 尝试匹配角色行模式
            patterns = [
                r'^(.*?)：男，(.+)$',  # 萧辞：男，双向救赎爱情线...
                r'^(.*?)：女，(.+)$',  # 棠月：女，双向救赎爱情线...
                r'^(.*?)：(.+)$',      # 通用格式：名字：描述
            ]
            
            matched = False
            for pattern in patterns:
                match = re.match(pattern, para)
                if match:
                    name_part = match.group(1).strip()
                    desc = match.group(2).strip()
                    
                    # 清理名字部分，移除引号等
                    name = name_part.replace('"', '').replace("'", "").strip()
                    
                    # 确定性别
                    gender = "male"
                    if pattern.startswith(r'^(.*?)：女，'):
                        gender = "female"
                    elif "适合" in desc and ("女性" in desc or "女生" in desc):
                        gender = "female"
                    elif "适合" in desc and ("男性" in desc or "男生" in desc):
                        gender = "male"
                    
                    # 过滤掉明显不是角色的行
                    if name not in ["适合", "DM", "NPC", "雷点", "主情感线", "主沉浸", "主情绪", "男生角色", "女生角色"]:
                        roles.append({
                            "name": name,
                            "gender": gender,
                            "desc": desc
                        })
                        matched = True
                        break
        
        # 去重（因为可能有重复）
        unique_roles = []
        seen_names = set()
        for role in roles:
            if role["name"] not in seen_names:
                unique_roles.append(role)
                seen_names.add(role["name"])
        
        print(f"提取角色数: {len(unique_roles)}")
        for i, role in enumerate(unique_roles[:10]):
            print(f"  {i+1}. {role['name']} ({role['gender']}): {role['desc'][:60]}...")
        
        return unique_roles
        
    except Exception as e:
        print(f"提取失败: {e}")
        return []

def create_update_data():
    """创建更新数据"""
    print("=== 创建网站更新数据 ===")
    
    # 首先从两个Word文档提取数据
    qingjun_path = "C:/Users/admin/Desktop/请君入瓮.docx"
    wanren_path = "C:/Users/admin/Desktop/万仞青山听水寒.docx"
    
    # 从请君入瓮.docx提取剧本简介
    script_descriptions = {}
    if os.path.exists(qingjun_path):
        script_descriptions = extract_script_descriptions(qingjun_path)
    else:
        print(f"文件不存在: {qingjun_path}")
    
    # 从万仞青山听水寒.docx提取角色信息
    wanren_roles = []
    if os.path.exists(wanren_path):
        wanren_roles = extract_wanren_roles(wanren_path)
    else:
        print(f"文件不存在: {wanren_path}")
    
    # 创建完整的更新数据
    all_scripts = ["请君入瓮", "流芳", "如故", "阙落", "万仞青山听水寒", "梦灵祈念时", "窃云台", "壁上观", "同归", "大宴之上"]
    
    update_data = {}
    
    for script_name in all_scripts:
        script_id = {
            "请君入瓮": "qingjunruweng",
            "万仞青山听水寒": "wanren", 
            "流芳": "liufang",
            "如故": "rugu",
            "阙落": "queluo",
            "梦灵祈念时": "mengling",
            "窃云台": "qieyuntai",
            "壁上观": "bishangguan",
            "同归": "tonggui",
            "大宴之上": "dayan"
        }.get(script_name, script_name.lower().replace(" ", ""))
        
        # 获取剧本简介
        description = ""
        if script_name in script_descriptions:
            description = script_descriptions[script_name]["description"]
        else:
            # 使用默认描述占位符
            description = f"{script_name}的详细剧情简介"
        
        # 获取角色信息（目前只有万仞青山听水寒有）
        roles = []
        if script_name == "万仞青山听水寒":
            # 只取前6个角色（通常每个剧本6个角色）
            roles = wanren_roles[:6]
        
        update_data[script_id] = {
            "script_name": script_name,
            "description": description,
            "roles": roles,
            "has_roles": len(roles) > 0,
            "button_should_exist": len(roles) > 0  # 如果有角色信息，显示按钮
        }
    
    print(f"\n=== 更新数据摘要 ===")
    for script_id, data in update_data.items():
        print(f"\n{data['script_name']} ({script_id}):")
        print(f"  简介长度: {len(data['description'])}")
        print(f"  角色数量: {len(data['roles'])}")
        print(f"  显示按钮: {'是' if data['button_should_exist'] else '否'}")
        if data['description'] and len(data['description']) < 100:
            print(f"  简介内容: {data['description']}")
    
    # 保存到文件
    output_file = "script_update_data.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(update_data, f, ensure_ascii=False, indent=2)
    
    print(f"\n更新数据已保存到: {output_file}")
    
    # 创建更新脚本
    create_update_script(update_data)
    
    return update_data

def create_update_script(update_data):
    """创建更新脚本"""
    print("\n=== 创建HTML更新脚本 ===")
    
    script_content = """#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
    
    update_examples = []
    
    for script_id, data in update_data.items():
        script_name = data['script_name']
        
        # 更新简介的例子
        if data['description']:
            old_desc_pattern = f'<p>{script_name}.*?</p>'
            new_desc = f'<p>{data["description"]}</p>'
            update_examples.append(f'# 更新 {script_name} 的简介')
            update_examples.append(f'replace_in_file(filePath, "{old_desc_pattern}", "{new_desc}", flags=re.DOTALL)')
        
        # 如果有角色信息，需要更新JavaScript中的roleData
        if data['has_roles']:
            update_examples.append(f'\n# 更新 {script_name} 的角色数据')
            
            roles_js = 'const roles = [\n'
            for role in data['roles']:
                roles_js += f'          {{ name: "{role["name"]}", gender: "{role["gender"]}", desc: "{role["desc"][:100]}" }},\n'
            roles_js += '        ]'
            
            update_examples.append(f'# 需要更新JavaScript中的{script_id}角色数据')
            update_examples.append(roles_js)
        
        # 如果不需要角色按钮，需要移除
        if not data['button_should_exist'] and script_name != "大宴之上":
            update_examples.append(f'\n# 移除 {script_name} 的角色简介按钮（如果存在）')
            update_examples.append(f'# 需要查找并移除 onclick="showRoleModal(\'{script_id}\')" 的按钮')
    
    script_content += "\n".join(update_examples)
    
    with open("update_website_instructions.txt", "w", encoding="utf-8") as f:
        f.write(script_content)
    
    print("更新指南已保存到: update_website_instructions.txt")
    
    return script_content

def main():
    """主函数"""
    print("开始分析剧本简介和角色信息...")
    
    # 获取当前目录
    current_dir = os.getcwd()
    print(f"工作目录: {current_dir}")
    
    # 创建更新数据
    update_data = create_update_data()
    
    # 显示总结
    print(f"\n{'='*60}")
    print("更新策略摘要:")
    print("="*60)
    
    total_scripts = len(update_data)
    scripts_with_roles = sum(1 for data in update_data.values() if data['has_roles'])
    scripts_with_desc = sum(1 for data in update_data.values() if data['description'] and len(data['description']) > 20)
    
    print(f"总剧本数: {total_scripts}")
    print(f"有角色信息的剧本: {scripts_with_roles}")
    print(f"有简介更新的剧本: {scripts_with_desc}")
    print(f"需要移除按钮的剧本: {total_scripts - scripts_with_roles}")
    
    print(f"\n具体操作:")
    print("1. 对于万仞青山听水寒: 更新角色信息，保留按钮")
    print("2. 对于其他剧本: 根据文档中的简介更新描述")
    print("3. 对于没有角色信息的剧本: 移除角色简介按钮")
    print("4. 对于大宴之上: 保留原按钮（如果网站已有）")
    
    return update_data

if __name__ == "__main__":
    result = main()